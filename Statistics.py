import pdb

#TODO Revisar
class MultiComparisons():
    """docstring forMultiANOVA2V."""
    import statsmodels.api as sm
    
    from statsmodels.formula.api import ols
    from pandas                  import DataFrame
    
    def __init__(self, data, columns, treatment, id, pvalue=0.001):
        self.data      = data
        self.columns   = columns
        self.treatment = treatment
        self.id        = id
        self.PVALUE    = pvalue
    
    def get_fullpairs(self):
        treatments = self.data[self.treatment].unique()
        
        paired = []
        for tr in treatments:
            query  = "{}=='{}'".format(self.treatment,tr)
            ids_tr = self.data.query(query)[self.id].unique()
            paired.append(set(ids_tr))
        
        paired = list(set.intersection(*paired))
        mask   = self.data[self.id].isin(paired)
        self.df = self.data[mask]
    
    def multiple_anova(self,formula,typ=2):
        Tests = self.DataFrame()
        for c in self.columns:
            #pdb.set_trace()
            
            self.get_fullpairs()
            model = self.ols(formula%c,self.df).fit()
            test  = self.sm.stats.anova_lm(model, typ=typ)
            
            test['analisis'] = c
            
            Tests = Tests.append(test)
        
        return Tests

#TODO Revisar
class Tables():
    """docstring for Tables."""

    def __init__(self, data, fname='results',pvalue=0.001):
        self.data   = data
        self.pvalue = pvalue
        self.fname  = fname
    
    def dt_to_docx(self):
        import docx
        
        from numpy import issubdtype, number, isnan
        
        shape = self.data.shape
        doc   = docx.Document()
        table = doc.add_table(shape[0]+1, shape[1])
        
        for i,d in enumerate(self.data.columns):
            table.cell(0,i).text = d
            run = table.cell(0,i).paragraphs[0].runs[0]
            run.font.bold = True
            
        for i,row in enumerate(self.data.values):
            for j,r in enumerate(row):
                test = issubdtype(type(r), number)
                if test:
                    text = '{:6.3f}'.format(r)
                    if r<=self.pvalue:
                        table.cell(i+1,j).text = '<{}'.format(self.pvalue)
                        run = table.cell(i+1,j).paragraphs[0].runs[0]
                        run.font.bold = True
                    elif isnan(r):
                        table.cell(i+1,j).text = '--'
                    else:
                        table.cell(i+1,j).text = text
                else:
                    table.cell(i+1,j).text = r
        doc.save('{}.docx'.format(self.fname))

class NormalTest():
    """docstring for ."""
    import pandas      as pd
    import numpy       as np
    import scipy.stats
    
    import statsmodels.stats.diagnostic as statsmodels
    
    def __init__(self, pvalue=0.05, n_opinions=1):
        self.test       = ['K-S','A-D','S-W']
        self.pvalue     = pvalue
        self.n_opinions = n_opinions
        #TODO: to implement n_opinios as a list for selecting which 
        #tests are being considered for analysis
        #TODO: Automated report, probablemente una clase aparte
        #self.pivot_index   = pivot_index
        #TODO: update this list in the automated report generation
        self.pivot_values  = ['mean','std','count','K-S','A-D','S-W','normality']
        self.pivot_columns = ['transformation']
        
    def __normal_opinion(self, df:np.array):
        """
        Counts how many pvalues in x rows are greater than pvalue and returns 'Gaussian' if 
        the account is greater than self.n_opinios; else the function returns 'No Gaussian'
        
        Parameters:
        ----------
        df : numpy array or pandas DataFrame -> (variables_tested, n_pvalues)
            Matrix of pvalues
        
        Returns:
        -------
          return :  dict
              - votes : number of test with p>n_pvalues for each variable tested
              - normality: Gaussian or No Gaussian if votes>n_pvalues
        """
        
        votes       = (df>self.pvalue).sum(axis=1)
        is_gaussian = votes>self.n_opinions
        is_gaussian = is_gaussian.replace({True:'Gaussian',False:'No Gaussian'})
        
        return {'votes':votes,'normality':is_gaussian}
    
    def norm_test(self, x : pd.Series):
        """
         Perform normality Tests
         
         Parameters:
         ----------
         x : pandas Series
         
         Returns:
         -------
         stats : dict
          - param : Series name
          - count : count
          - media : mean
          - std   : standard deviation
          - K-S   : Kolmogorov-Smirnov
          - A-D   : Anderson-Darling
          - S-K   : Shapiro-Wilk
        """
        #TODO : automatizar cuales pruebas se van a realizar -> llamadas a funciones tipo {'x':func}
        stats = { 'param' :x.name  ,
                  'count' :x.count(),
                  'mean'  :x.mean (),
                  'std'   :x.std  (),
                  #'K-S'  :self.st.kstest (x,'norm')[1],
                  'K-S'   :self.statsmodels.kstest_fit(x)[1],
                  'A-D'   :self.statsmodels.normal_ad (x)[1],
                  'S-W'   :self.scipy.stats.shapiro   (x)[1]  }

        return stats
    
    def norm_test_df(self, df:pd.DataFrame):
        """
        Perform nomal tests (using self.norm_test) over a numerical DataFrame
        
        Parameters:
        ----------
        df : Pandas DataFrame
            Numerical DataFrame
                    
        Returns:
        -------
          table : DataFrame
            Rows:
              A row for each column in df
            
            Columns:
              - param : Series name
              - count : count
              - media : mean
              - std   : standard deviation
              - K-S   : Kolmogorov-Smirnov
              - A-D   : Anderson-Darling
              - S-K   : Shapiro-Wilk
        """
        #TODO : automatizar cuales pruebas se van a realizar -> llamadas a funciones tipo {'x':func}
        table = self.pd.DataFrame(columns=['param','count','mean','std','K-S','A-D','S-W'])
        for col in df.columns:
            tab   = self.norm_test(df[col])
            table = table.append(tab,ignore_index=True)
        
        return table
    
    def norm_test_df_groups(self,df,groups):
        """
        Perform normal test over the numerical columns in DataFrame df,
        considering gruped data.
        These analysis also include a log10 transformation of df
        
        Parameters:
        ----------
        df : pandas DataFrame (variables_tested, )
             DataFrame get from the evaluation of self.norm_test_df
        groups : list
            List of grouping
        
        Returns:
        -------
        return : type
            return description
        """
        #TODO: generalize for k transformations
        table = self.pd.DataFrame()
        for grpnm,dfgrp in df.groupby(groups):
            to_test    = dfgrp.select_dtypes('number')
            tmp        = self.norm_test_df(to_test)
            if type(groups)==list:
                treatments = {g:i for g,i in zip(groups, grpnm)}
            else:
                treatments = {groups:grpnm}
            
            treatments   .update({'transformation':'identity'})
            tmp        = tmp.assign(**treatments)
            table      = table.append(tmp)
            
            tmp        = self.norm_test_df(self.np.log10(to_test))
            treatments   .update({'transformation':'log10'})
            tmp        = tmp.assign(**treatments)
            table      = table.append(tmp)
                
        table        = table.reset_index(drop=True)
        norm_opinion = self.__normal_opinion(table.loc[:,self.test])
        table        = table.assign(**norm_opinion)
        #pdb.set_trace()
        #TODO: generalize for k transformations
        #table = table.pivot(index=self.pivot_index,columns=self.pivot_columns,values=self.pivot_values)
        #self.pd.options.display.float_format = '{:,.3f}'.format
        #table = {'identity':table.xs('identity', level='transformation',axis=1),
        #         'log10'   :table.xs('log10'   , level='transformation',axis=1) }
        
        return table

class TwoTreatmentTests():
    """docstring for ."""
    import pandas      as pd
    import numpy       as np
    
    import pingouin    as pg
    
    def __init__(self, pvalue=0.001,params={'paired':False,'tail':'two-sided'}):
        self.test          = []
        self.pvalue        = pvalue
        self.params        = params
        self.pivot_values  = ['mean','std','count','K-S','A-D','S-W','normality']
        self.pivot_columns = ['transformation']
    
    def ttest(self, x, y, params={}):
        """
        Perform T-Student test wrapping pingouin.ttest 
        
        Parameters:
        ----------
        x : list, numpy (samples), pandas.Series
            observations
        y : list, numpy (samples), pandas.Series, int, float
            observations or single value
        
        Returns:
        -------
         results : DataFrame
          Results T-Student Tests
        """
        if bool(params):
            params = self.params
        
        result = self.pg.ttest( x,y,**params)
        return result
    
    #TODO: paired ttest implementation
    def df_ttest(self, x:pd.DataFrame, group:str, ttype:str = 'independent'):
        """
         Perform T-Student tests over a long format DataFrame. If the grouping column
         has not exactly two categories then a ValueError exception is thrown.
         
         Parameters:
         ----------
         x : DataFrame in long format
         group : string
           Name of the grouping column with only 2 categories
         ttype: string
           Indicates paired or not paired T-Test NOTE: not yet implemented
         
         Returns:
         -------
         results : DataFrame
          Results of all the T-Student Tests
        """
        groups  = list(x[group].unique())
        columns = x.select_dtypes('number').columns
        table   = self.pd.DataFrame()
        if len(groups)!=2:
            raise ValueError('Number of groups is not 2')
        else:
            df0 = x[x[group]==groups[0]]
            df1 = x[x[group]==groups[1]]
            for col in columns:
                tmp   = self.ttest( df0[col],df1[col] )
                tmp   = tmp.assign(**{'feature':col})
                tmp   = tmp.assign(**{'%s mean'%groups[0]:df0[col].mean(),
                                      '%s std' %groups[0]:df0[col].std (),
                                      '%s mean'%groups[1]:df1[col].mean(),
                                      '%s std' %groups[1]:df1[col].std () })
                table = table.append(tmp)
            
            #TODO: make non-parametric tests
            """
            from scipy.stats import mannwhitneyu as utest
            col = 'IMC Preg'
            utest( df0[col],df1[col] )
                
            pdb.set_trace()
            """
        return table
        
    #TODO: paired ttest implementation
    def df_full_ttest(self, x:pd.DataFrame, group:str, ttype:str = 'independent'):
        table              = self.df_ttest(x,group)
        table['transform'] = 'raw'
        
        df       = x.select_dtypes(include='number')
        non_zero = (df<=0).sum().values
        non_zero = df.columns[non_zero==0].to_list()
        if len(non_zero)>0:
            df_transform     = self.np.log10(x[non_zero])
            tmp              = self.df_ttest( df_transform.join(x[group]), group )
            tmp['transform'] = 'log10'
            table            = table.append(tmp)
        else:
            print("No features log10 transformed due to some values are less than 0")
        
        return table